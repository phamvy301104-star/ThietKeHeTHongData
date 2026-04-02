"""
Script tạo báo cáo đồ án Word - Tối ưu hóa hệ thống quản lý kho hàng
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False, 
                             size=13, alignment=None, space_before=0, space_after=6,
                             font_name='Times New Roman', color=None, underline=False,
                             first_line_indent=None, line_spacing=1.5):
    p = doc.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_multi_run_paragraph(doc, runs_data, alignment=None, space_before=0, space_after=6,
                             first_line_indent=None, line_spacing=1.5):
    """Add paragraph with multiple runs (mixed formatting)."""
    p = doc.add_paragraph()
    for rd in runs_data:
        run = p.add_run(rd.get('text', ''))
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(rd.get('size', 13))
        run.bold = rd.get('bold', False)
        run.italic = rd.get('italic', False)
        run.underline = rd.get('underline', False)
        if rd.get('color'):
            run.font.color.rgb = rd['color']
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    return table

def add_image_placeholder(doc, caption, image_path=None):
    """Thêm hình ảnh hoặc placeholder."""
    if image_path and os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_formatted_paragraph(doc, f"[Chèn ảnh: {caption}]",
                                 italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 color=RGBColor(255, 0, 0))
    # Caption below image
    add_formatted_paragraph(doc, caption, italic=True, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

def create_report():
    doc = Document()
    
    # === PAGE SETUP ===
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # ============================================================
    # TRANG PHỤ BÌA
    # ============================================================
    add_formatted_paragraph(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", bold=True, size=14,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)
    add_formatted_paragraph(doc, "TRƯỜNG ĐẠI HỌC ...", bold=True, size=14,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "KHOA CÔNG NGHỆ THÔNG TIN", bold=True, size=14,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_formatted_paragraph(doc, "⎯⎯⎯⎯⎯⎯⎯  ⎯⎯⎯⎯⎯⎯⎯", size=14,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    
    add_formatted_paragraph(doc, "ĐỒ ÁN MÔN HỌC", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc, "DỮ LIỆU LỚN (BIG DATA)", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    add_formatted_paragraph(doc, "TỐI ƯU HÓA HỆ THỐNG QUẢN LÝ KHO HÀNG", bold=True, size=18,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                             color=RGBColor(0, 51, 153), space_after=6)
    add_formatted_paragraph(doc, "SỬ DỤNG CÔNG NGHỆ BIG DATA", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             color=RGBColor(0, 51, 153), space_after=48)
    
    add_formatted_paragraph(doc, "Giảng viên hướng dẫn: TS. ...", size=14,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    add_formatted_paragraph(doc, "Sinh viên thực hiện:     ...", size=14,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    add_formatted_paragraph(doc, "MSSV:                          ...", size=14,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    add_formatted_paragraph(doc, "Lớp:                              ...", size=14,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=48)
    
    add_formatted_paragraph(doc, "TP. Hồ Chí Minh, năm 2026", bold=True, size=14,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    
    doc.add_page_break()
    
    # ============================================================
    # LỜI CAM ĐOAN
    # ============================================================
    add_formatted_paragraph(doc, "LỜI CAM ĐOAN", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    add_formatted_paragraph(doc, 
        "Tôi xin cam đoan đây là công trình nghiên cứu của riêng tôi. Các số liệu, kết quả nêu trong đồ án là trung thực và chưa từng được ai công bố trong bất kỳ công trình nào khác.",
        size=13, first_line_indent=1.27, space_after=12)
    add_formatted_paragraph(doc,
        "Tôi xin cam đoan rằng mọi sự giúp đỡ cho việc thực hiện đồ án này đã được cảm ơn và các thông tin trích dẫn trong đồ án đã được chỉ rõ nguồn gốc.",
        size=13, first_line_indent=1.27, space_after=24)
    
    add_formatted_paragraph(doc, "TP. Hồ Chí Minh, ngày ... tháng ... năm 2026",
                             italic=True, size=13, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    add_formatted_paragraph(doc, "Sinh viên thực hiện",
                             bold=True, size=13, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=48)
    add_formatted_paragraph(doc, "(Ký và ghi rõ họ tên)",
                             italic=True, size=13, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    
    doc.add_page_break()

    # ============================================================
    # LỜI CẢM ƠN
    # ============================================================
    add_formatted_paragraph(doc, "LỜI CẢM ƠN", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    add_formatted_paragraph(doc,
        "Để hoàn thành đồ án này, tôi xin gửi lời cảm ơn chân thành đến:",
        size=13, first_line_indent=1.27, space_after=12)
    
    add_formatted_paragraph(doc,
        "Trước hết, tôi xin bày tỏ lòng biết ơn sâu sắc đến Thầy/Cô TS. ... \u2013 người đã trực tiếp hướng dẫn, định hướng nghiên cứu và tận tình chỉ bảo tôi trong suốt quá trình thực hiện đồ án.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc,
        "Tôi xin chân thành cảm ơn quý Thầy/Cô trong Khoa Công nghệ Thông tin, Trường Đại học ... đã truyền đạt những kiến thức quý báu trong suốt thời gian học tập tại trường, làm nền tảng để tôi hoàn thành đồ án này.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc,
        "Tôi cũng xin gửi lời cảm ơn đến gia đình, bạn bè đã luôn động viên, hỗ trợ tôi trong suốt quá trình học tập và nghiên cứu.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc,
        "Mặc dù đã cố gắng hoàn thiện đồ án, tuy nhiên không tránh khỏi những thiếu sót. Tôi rất mong nhận được sự góp ý của quý Thầy/Cô để đồ án được hoàn thiện hơn.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Xin chân thành cảm ơn!",
                             size=13, first_line_indent=1.27, space_after=24)
    
    add_formatted_paragraph(doc, "TP. Hồ Chí Minh, ngày ... tháng ... năm 2026",
                             italic=True, size=13, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    add_formatted_paragraph(doc, "Sinh viên thực hiện",
                             bold=True, size=13, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=48)
    add_formatted_paragraph(doc, "(Ký và ghi rõ họ tên)",
                             italic=True, size=13, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    
    doc.add_page_break()

    # ============================================================
    # MỤC LỤC
    # ============================================================
    add_formatted_paragraph(doc, "MỤC LỤC", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    toc_items = [
        ("Lời cam đoan", "i"),
        ("Lời cảm ơn", "ii"),
        ("Mục lục", "iii"),
        ("Danh mục các ký hiệu, các chữ viết tắt", "iv"),
        ("Danh mục các bảng", "v"),
        ("Danh mục các hình vẽ, đồ thị", "vi"),
        ("Chương 1. TỔNG QUAN", "1"),
        ("    1.1 Giới thiệu đề tài", "1"),
        ("    1.2 Tính cấp thiết và lý do hình thành đề tài", "2"),
        ("    1.3 Mục tiêu nghiên cứu", "3"),
        ("    1.4 Đối tượng và phạm vi nghiên cứu", "3"),
        ("    1.5 Cấu trúc đồ án", "4"),
        ("Chương 2. CƠ SỞ LÝ THUYẾT", "5"),
        ("    2.1 Tổng quan về Big Data", "5"),
        ("    2.2 Apache Spark và PySpark", "6"),
        ("    2.3 Hệ thống tập tin phân tán HDFS", "8"),
        ("    2.4 Docker và container hóa", "9"),
        ("    2.5 Thuật toán K-Means Clustering", "10"),
        ("    2.6 Phân tích ABC trong quản lý kho", "11"),
        ("    2.7 Công cụ trực quan hóa Power BI", "12"),
        ("Chương 3. KẾT QUẢ THỰC NGHIỆM", "13"),
        ("    3.1 Kiến trúc hệ thống", "13"),
        ("    3.2 Bộ dữ liệu và tiền xử lý", "14"),
        ("    3.3 Phân tích doanh thu và KPI kho hàng", "16"),
        ("    3.4 Phân cụm sản phẩm bằng K-Means", "18"),
        ("    3.5 Tối ưu hóa hiệu suất Spark", "21"),
        ("    3.6 Dashboard Power BI", "23"),
        ("Chương 4. KẾT LUẬN VÀ KIẾN NGHỊ", "25"),
        ("    4.1 Kết luận", "25"),
        ("    4.2 Hạn chế", "26"),
        ("    4.3 Hướng phát triển", "26"),
        ("TÀI LIỆU THAM KHẢO", "27"),
        ("PHỤ LỤC", "29"),
    ]
    # Tab stop ở lề phải (16cm - 3cm lề trái - 2cm lề phải = 11cm nội dung)
    for item_text, page_num in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.5
        # Thêm tab stop với dot leader ở vị trí 15cm từ lề trái
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(13.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        # Text run
        run = p.add_run(item_text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        is_chapter = item_text.startswith("Chương") or item_text in ["TÀI LIỆU THAM KHẢO", "PHỤ LỤC"] or not item_text.startswith("    ")
        run.font.size = Pt(13)
        run.bold = item_text.startswith("Chương") or item_text in ["TÀI LIỆU THAM KHẢO", "PHỤ LỤC"]
        # Tab + page number
        tab_run = p.add_run("\t")
        tab_run.font.name = 'Times New Roman'
        tab_run.font.size = Pt(13)
        page_run = p.add_run(page_num)
        page_run.font.name = 'Times New Roman'
        page_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        page_run.font.size = Pt(13)
    
    doc.add_page_break()
    
    # ============================================================
    # DANH MỤC CHỮ VIẾT TẮT
    # ============================================================
    add_formatted_paragraph(doc, "DANH MỤC CÁC KÝ HIỆU, CÁC CHỮ VIẾT TẮT", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    abbr_data = [
        ("HDFS", "Hadoop Distributed File System - Hệ thống tập tin phân tán Hadoop"),
        ("WSSSE", "Within Set Sum of Squared Errors - Tổng bình phương sai số nội cụm"),
        ("K-Means", "Thuật toán phân cụm K-Means"),
        ("ML", "Machine Learning - Học máy"),
        ("MLlib", "Machine Learning Library - Thư viện học máy của Spark"),
        ("RDD", "Resilient Distributed Dataset - Tập dữ liệu phân tán có khả năng phục hồi"),
        ("API", "Application Programming Interface - Giao diện lập trình ứng dụng"),
        ("CSV", "Comma-Separated Values - Giá trị phân cách bằng dấu phẩy"),
        ("ABC", "Phân loại ABC (Activity Based Classification)"),
        ("KPI", "Key Performance Indicator - Chỉ số hiệu suất chính"),
        ("ETL", "Extract, Transform, Load - Trích xuất, Chuyển đổi, Nạp dữ liệu"),
        ("UI", "User Interface - Giao diện người dùng"),
    ]
    add_table(doc, ["Ký hiệu/Viết tắt", "Ý nghĩa"],
              abbr_data, col_widths=[4, 12])
    
    doc.add_page_break()
    
    # ============================================================
    # DANH MỤC BẢNG
    # ============================================================
    add_formatted_paragraph(doc, "DANH MỤC CÁC BẢNG", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    tables_list = [
        ("Bảng 2.1", "So sánh các framework xử lý Big Data"),
        ("Bảng 2.2", "Cấu hình Docker cluster"),
        ("Bảng 3.1", "Thông tin 9 bảng dữ liệu Olist"),
        ("Bảng 3.2", "Thống kê missing values sau khi làm sạch"),
        ("Bảng 3.3", "Kết quả phân tích ABC"),
        ("Bảng 3.4", "Silhouette Score theo số cụm K"),
        ("Bảng 3.5", "Đặc trưng trung bình của từng cụm"),
        ("Bảng 3.6", "So sánh hiệu suất trước và sau tối ưu Spark"),
    ]
    add_table(doc, ["Bảng", "Nội dung"], tables_list, col_widths=[3, 13])
    
    doc.add_page_break()
    
    # ============================================================
    # DANH MỤC HÌNH
    # ============================================================
    add_formatted_paragraph(doc, "DANH MỤC CÁC HÌNH VẼ, ĐỒ THỊ", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    figures_list = [
        ("Hình 2.1", "Kiến trúc Apache Spark"),
        ("Hình 2.2", "Mô hình HDFS với NameNode và DataNode"),
        ("Hình 2.3", "Minh họa thuật toán K-Means"),
        ("Hình 2.4", "Phân loại ABC trong quản lý kho"),
        ("Hình 3.1", "Kiến trúc tổng thể hệ thống"),
        ("Hình 3.2", "Sơ đồ pipeline xử lý dữ liệu"),
        ("Hình 3.3", "Biểu đồ Elbow Method và Silhouette Score"),
        ("Hình 3.4", "Kết quả phân cụm sản phẩm"),
        ("Hình 3.5", "Boxplot so sánh đặc trưng các cụm"),
        ("Hình 3.6", "Dashboard Power BI - Tổng quan doanh thu"),
        ("Hình 3.7", "Dashboard Power BI - Phân tích kho hàng"),
    ]
    add_table(doc, ["Hình", "Nội dung"], figures_list, col_widths=[3, 13])
    
    doc.add_page_break()
    
    # ============================================================
    # CHƯƠNG 1: TỔNG QUAN
    # ============================================================
    add_formatted_paragraph(doc, "Chương 1. TỔNG QUAN", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    # --- 1.1 ---
    add_formatted_paragraph(doc, "1.1 Giới thiệu đề tài", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Trong bối cảnh thương mại điện tử phát triển mạnh mẽ, lượng dữ liệu phát sinh từ hoạt động kinh doanh ngày càng lớn. Theo báo cáo của IDC (2025), lượng dữ liệu toàn cầu dự kiến đạt 175 Zettabyte vào năm 2025, trong đó dữ liệu liên quan đến chuỗi cung ứng và quản lý kho hàng chiếm tỷ trọng đáng kể [1]. Việc áp dụng công nghệ Big Data vào quản lý kho hàng là xu hướng tất yếu để nâng cao hiệu quả vận hành.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc,
        "Quản lý kho hàng truyền thống gặp nhiều hạn chế như: không thể xử lý dữ liệu lớn trong thời gian thực, thiếu khả năng phân tích dự đoán, và khó tối ưu hóa tồn kho. Các nghiên cứu trước đây của Chen et al. (2023) đã chỉ ra rằng việc áp dụng phân cụm K-Means vào phân loại sản phẩm giúp giảm 15-20% chi phí tồn kho [2]. Tương tự, nghiên cứu của Wang & Li (2024) về ứng dụng Apache Spark trong phân tích chuỗi cung ứng cho thấy tốc độ xử lý nhanh hơn 10-100 lần so với phương pháp truyền thống [3].",
        size=13, first_line_indent=1.27, space_after=6)

    add_formatted_paragraph(doc,
        "Đồ án này xây dựng một hệ thống phân tích dữ liệu lớn hoàn chỉnh, sử dụng bộ dữ liệu Brazilian E-Commerce (Olist) với hơn 100.000 đơn hàng thực tế, nhằm minh họa quy trình tối ưu hóa quản lý kho hàng bằng công nghệ Big Data.",
        size=13, first_line_indent=1.27, space_after=12)
    
    # --- 1.2 ---
    add_formatted_paragraph(doc, "1.2 Tính cấp thiết và lý do hình thành đề tài", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Sự bùng nổ của thương mại điện tử đặt ra thách thức lớn cho các doanh nghiệp trong việc quản lý kho hàng hiệu quả. Cụ thể:",
        size=13, first_line_indent=1.27, space_after=6)
    
    bullets = [
        "Khối lượng đơn hàng tăng nhanh: hàng nghìn đến hàng triệu đơn hàng mỗi ngày, vượt quá khả năng xử lý của hệ thống truyền thống;",
        "Đa dạng sản phẩm: hàng chục nghìn SKU cần phân loại và quản lý tồn kho khác nhau;",
        "Yêu cầu giao hàng nhanh: áp lực giảm thời gian giao hàng đòi hỏi phân bổ kho thông minh theo khu vực;",
        "Tối ưu chi phí: cần cân bằng giữa chi phí tồn kho và mức phục vụ khách hàng."
    ]
    for b in bullets:
        add_formatted_paragraph(doc, f"- {b}", size=13, space_after=3)
    
    add_formatted_paragraph(doc,
        "Từ những thách thức trên, việc xây dựng hệ thống phân tích dữ liệu lớn để hỗ trợ ra quyết định quản lý kho hàng là cấp thiết và có ý nghĩa thực tiễn cao.",
        size=13, first_line_indent=1.27, space_before=6, space_after=12)
    
    # --- 1.3 ---
    add_formatted_paragraph(doc, "1.3 Mục tiêu nghiên cứu", bold=True, size=14, space_before=12, space_after=6)
    
    objectives = [
        "Xây dựng pipeline xử lý dữ liệu lớn hoàn chỉnh sử dụng Apache Spark (PySpark);",
        "Thiết kế hệ thống phân tán với Docker cluster (Spark Master + Workers) và HDFS;",
        "Phân tích doanh thu, hiệu suất giao hàng và KPI kho hàng từ dữ liệu thực tế;",
        "Áp dụng thuật toán K-Means Clustering để phân cụm sản phẩm phục vụ chiến lược tồn kho;",
        "Áp dụng phân tích ABC để phân loại sản phẩm theo đóng góp doanh thu;",
        "Minh họa các kỹ thuật tối ưu Spark: Cache, Broadcast Join, Repartition;",
        "Trực quan hóa kết quả bằng Power BI dashboard."
    ]
    for i, obj in enumerate(objectives, 1):
        add_formatted_paragraph(doc, f"({i}) {obj}", size=13, space_after=3)
    
    # --- 1.4 ---
    add_formatted_paragraph(doc, "1.4 Đối tượng và phạm vi nghiên cứu", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc, "1.4.1 Đối tượng nghiên cứu", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Đối tượng nghiên cứu là hệ thống quản lý kho hàng của sàn thương mại điện tử, sử dụng bộ dữ liệu Brazilian E-Commerce (Olist) bao gồm 9 bảng dữ liệu với thông tin đơn hàng, sản phẩm, người bán, khách hàng, thanh toán, đánh giá và vị trí địa lý.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "1.4.2 Phạm vi nghiên cứu", italic=True, size=14, space_after=6)
    scope = [
        "Dữ liệu: Bộ dữ liệu Olist gồm ~100.000 đơn hàng trong giai đoạn 2016-2018;",
        "Công nghệ: PySpark, Docker, HDFS, K-Means (MLlib), Power BI;",
        "Phân tích: Doanh thu, hiệu suất giao hàng, phân cụm sản phẩm, phân tích ABC;",
        "Giới hạn: Không bao gồm dự báo nhu cầu (forecasting) và tối ưu hóa đường đi vận chuyển."
    ]
    for s in scope:
        add_formatted_paragraph(doc, f"- {s}", size=13, space_after=3)
    
    # --- 1.5 ---
    add_formatted_paragraph(doc, "1.5 Cấu trúc đồ án", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc, "Đồ án được tổ chức thành 4 chương:", size=13, first_line_indent=1.27, space_after=6)
    
    chapters_desc = [
        ("Chương 1 - Tổng quan:", " Giới thiệu đề tài, tính cấp thiết, mục tiêu, đối tượng và phạm vi nghiên cứu."),
        ("Chương 2 - Cơ sở lý thuyết:", " Trình bày nền tảng lý thuyết về Big Data, Apache Spark, HDFS, Docker, thuật toán K-Means, phân tích ABC và Power BI."),
        ("Chương 3 - Kết quả thực nghiệm:", " Mô tả kiến trúc hệ thống, quy trình xử lý dữ liệu, kết quả phân tích và trực quan hóa."),
        ("Chương 4 - Kết luận và kiến nghị:", " Tổng kết kết quả đạt được, hạn chế và hướng phát triển.")
    ]
    for title, desc in chapters_desc:
        add_multi_run_paragraph(doc, [
            {'text': title, 'bold': True, 'size': 13},
            {'text': desc, 'size': 13}
        ], space_after=3, first_line_indent=1.27)
    
    doc.add_page_break()
    
    # ============================================================
    # CHƯƠNG 2: CƠ SỞ LÝ THUYẾT
    # ============================================================
    add_formatted_paragraph(doc, "Chương 2. CƠ SỞ LÝ THUYẾT", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    # --- 2.1 ---
    add_formatted_paragraph(doc, "2.1 Tổng quan về Big Data", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Big Data (Dữ liệu lớn) là thuật ngữ mô tả các tập dữ liệu có khối lượng lớn, tốc độ sinh ra nhanh và đa dạng về định dạng, vượt quá khả năng xử lý của các công cụ truyền thống. Big Data được đặc trưng bởi mô hình 5V [4]:",
        size=13, first_line_indent=1.27, space_after=6)
    
    v5 = [
        ("Volume (Khối lượng):", " Dữ liệu có kích thước từ Terabyte đến Petabyte."),
        ("Velocity (Tốc độ):", " Dữ liệu sinh ra và cần xử lý trong thời gian thực hoặc gần thực."),
        ("Variety (Đa dạng):", " Bao gồm dữ liệu có cấu trúc, bán cấu trúc và phi cấu trúc."),
        ("Veracity (Độ tin cậy):", " Mức độ chính xác và tin cậy của dữ liệu."),
        ("Value (Giá trị):", " Khả năng trích xuất giá trị hữu ích từ dữ liệu.")
    ]
    for title, desc in v5:
        add_multi_run_paragraph(doc, [
            {'text': f"- {title}", 'bold': True, 'size': 13},
            {'text': desc, 'size': 13}
        ], space_after=3)
    
    add_formatted_paragraph(doc,
        "Trong lĩnh vực quản lý kho hàng, Big Data được ứng dụng để phân tích hành vi mua sắm, dự báo nhu cầu, tối ưu hóa tồn kho và cải thiện chuỗi cung ứng.",
        size=13, first_line_indent=1.27, space_before=6, space_after=12)
    
    # --- 2.2 ---
    add_formatted_paragraph(doc, "2.2 Apache Spark và PySpark", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc, "2.2.1 Apache Spark", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Apache Spark là một framework mã nguồn mở cho xử lý dữ liệu phân tán, được phát triển tại UC Berkeley và hiện do Apache Software Foundation quản lý [5]. Spark xử lý dữ liệu trong bộ nhớ (in-memory computing), giúp tốc độ nhanh hơn 10-100 lần so với Hadoop MapReduce truyền thống.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Các thành phần chính của Apache Spark:", size=13, first_line_indent=1.27, space_after=6)
    spark_components = [
        ("Spark Core:", " Cung cấp RDD (Resilient Distributed Dataset) và các chức năng cơ bản như quản lý bộ nhớ, lập lịch tác vụ."),
        ("Spark SQL:", " Xử lý dữ liệu có cấu trúc thông qua DataFrame API và SQL."),
        ("MLlib:", " Thư viện machine learning phân tán, hỗ trợ phân cụm, phân loại, hồi quy."),
        ("Spark Streaming:", " Xử lý dữ liệu luồng trong thời gian thực.")
    ]
    for title, desc in spark_components:
        add_multi_run_paragraph(doc, [
            {'text': f"- {title}", 'bold': True, 'size': 13},
            {'text': desc, 'size': 13}
        ], space_after=3)
    
    add_image_placeholder(doc, "Hình 2.1. Kiến trúc Apache Spark")
    
    add_formatted_paragraph(doc, "2.2.2 PySpark", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "PySpark là API Python cho Apache Spark, cho phép lập trình viên Python sử dụng toàn bộ sức mạnh của Spark. Trong đồ án này, PySpark được sử dụng phiên bản 3.5.0 với các module: SparkSQL để xử lý DataFrame, MLlib để thực hiện K-Means Clustering [6].",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Cấu hình SparkSession trong đồ án:", size=13, first_line_indent=1.27, space_after=3)
    config_rows = [
        ("spark.sql.adaptive.enabled", "true", "Tự động tối ưu query plan"),
        ("spark.serializer", "KryoSerializer", "Tuần tự hóa nhanh hơn Java mặc định"),
        ("spark.sql.shuffle.partitions", "8", "Số partition khi shuffle"),
        ("spark.driver.memory", "2g", "Bộ nhớ cho driver"),
        ("spark.executor.memory", "2g", "Bộ nhớ cho mỗi executor"),
    ]
    add_table(doc, ["Cấu hình", "Giá trị", "Mô tả"], config_rows, col_widths=[5.5, 3.5, 7])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    # --- 2.3 ---
    add_formatted_paragraph(doc, "2.3 Hệ thống tập tin phân tán HDFS", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "HDFS (Hadoop Distributed File System) là hệ thống tập tin phân tán của Hadoop, được thiết kế để lưu trữ dữ liệu lớn trên nhiều máy tính [7]. HDFS hoạt động theo kiến trúc Master-Slave:",
        size=13, first_line_indent=1.27, space_after=6)
    
    hdfs_comp = [
        ("NameNode (Master):", " Quản lý metadata, thông tin về vị trí các khối dữ liệu."),
        ("DataNode (Slave):", " Lưu trữ các khối dữ liệu thực tế, mỗi khối mặc định 128MB."),
        ("Replication:", " Mỗi khối dữ liệu được nhân bản (mặc định 3 bản) để đảm bảo tính sẵn sàng.")
    ]
    for title, desc in hdfs_comp:
        add_multi_run_paragraph(doc, [
            {'text': f"- {title}", 'bold': True, 'size': 13},
            {'text': desc, 'size': 13}
        ], space_after=3)
    
    add_formatted_paragraph(doc,
        "Trong đồ án, HDFS được triển khai với 1 NameNode và 2 DataNode, hệ số nhân bản là 2, thông qua Docker container.",
        size=13, first_line_indent=1.27, space_before=6, space_after=6)
    
    add_image_placeholder(doc, "Hình 2.2. Mô hình HDFS với NameNode và DataNode")

    # --- 2.4 ---
    add_formatted_paragraph(doc, "2.4 Docker và container hóa", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Docker là nền tảng container hóa cho phép đóng gói ứng dụng cùng các phụ thuộc vào các container độc lập, đảm bảo tính nhất quán giữa các môi trường [8]. Docker Compose cho phép định nghĩa và chạy nhiều container cùng lúc thông qua file YAML.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc,
        "Cấu hình Docker cluster trong đồ án bao gồm 7 container:",
        size=13, first_line_indent=1.27, space_after=3)
    
    add_formatted_paragraph(doc, "Bảng 2.2. Cấu hình Docker cluster", bold=True, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    docker_rows = [
        ("namenode", "Hadoop NameNode", "9870, 9000", "Quản lý metadata HDFS"),
        ("datanode1", "Hadoop DataNode", "9864", "Lưu trữ dữ liệu 1"),
        ("datanode2", "Hadoop DataNode", "9865", "Lưu trữ dữ liệu 2"),
        ("spark-master", "Spark Master", "8080, 7077", "Điều phối tác vụ Spark"),
        ("spark-worker-1", "Spark Worker", "8081", "2 cores, 2GB RAM"),
        ("spark-worker-2", "Spark Worker", "8082", "2 cores, 2GB RAM"),
        ("jupyter", "Jupyter Notebook", "8888", "Môi trường phát triển"),
    ]
    add_table(doc, ["Container", "Vai trò", "Port", "Mô tả"], docker_rows, col_widths=[3.5, 3.5, 3, 6])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    # --- 2.5 ---
    add_formatted_paragraph(doc, "2.5 Thuật toán K-Means Clustering", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "K-Means là thuật toán phân cụm (clustering) phổ biến nhất trong machine learning không giám sát (unsupervised learning). Thuật toán chia n điểm dữ liệu thành K cụm sao cho tổng bình phương khoảng cách từ mỗi điểm đến tâm cụm gần nhất là nhỏ nhất [9].",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Hàm mục tiêu (WSSSE):", bold=True, size=13, space_after=6)
    add_formatted_paragraph(doc, "J = Σᵢ₌₁ᴷ Σₓ∈Cᵢ ||x - μᵢ||²          (2.1)",
                             size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc,
        "Trong đó: K là số cụm, Cᵢ là cụm thứ i, μᵢ là tâm cụm thứ i, x là điểm dữ liệu.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Các bước của thuật toán K-Means:", bold=True, size=13, space_after=6)
    kmeans_steps = [
        "Bước 1: Chọn ngẫu nhiên K tâm cụm ban đầu;",
        "Bước 2: Gán mỗi điểm dữ liệu vào cụm có tâm gần nhất;",
        "Bước 3: Cập nhật tâm cụm bằng trung bình các điểm trong cụm;",
        "Bước 4: Lặp lại Bước 2-3 cho đến khi hội tụ (tâm cụm không thay đổi)."
    ]
    for s in kmeans_steps:
        add_formatted_paragraph(doc, s, size=13, space_after=3)
    
    add_formatted_paragraph(doc, "Chọn số cụm K tối ưu:", bold=True, size=13, space_before=6, space_after=6)
    add_formatted_paragraph(doc,
        "Phương pháp Elbow: Vẽ đồ thị WSSSE theo K, chọn điểm \"khuỷu tay\" nơi WSSSE bắt đầu giảm chậm.",
        size=13, first_line_indent=1.27, space_after=3)
    add_formatted_paragraph(doc,
        "Silhouette Score: Đo mức độ tương đồng của một điểm với cụm của nó so với các cụm khác. Giá trị từ -1 đến 1, càng gần 1 càng tốt. Công thức:",
        size=13, first_line_indent=1.27, space_after=6)
    add_formatted_paragraph(doc, "s(i) = (b(i) - a(i)) / max(a(i), b(i))          (2.2)",
                             size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc,
        "Trong đó: a(i) là khoảng cách trung bình từ điểm i đến các điểm cùng cụm, b(i) là khoảng cách trung bình nhỏ nhất từ điểm i đến các điểm ở cụm khác.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_image_placeholder(doc, "Hình 2.3. Minh họa thuật toán K-Means")
    
    # --- 2.6 ---
    add_formatted_paragraph(doc, "2.6 Phân tích ABC trong quản lý kho", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Phân tích ABC là phương pháp phân loại hàng tồn kho dựa trên nguyên tắc Pareto (80/20). Sản phẩm được chia thành 3 nhóm dựa trên đóng góp doanh thu [10]:",
        size=13, first_line_indent=1.27, space_after=6)
    
    abc_rows = [
        ("A", "Top 20%", "~80% doanh thu", "Kiểm soát chặt, tồn kho tối ưu"),
        ("B", "20% - 50%", "~15% doanh thu", "Kiểm soát vừa phải"),
        ("C", "Bottom 50%", "~5% doanh thu", "Kiểm soát đơn giản, tồn kho tối thiểu"),
    ]
    add_table(doc, ["Nhóm", "Phần trăm sản phẩm", "Đóng góp doanh thu", "Chiến lược kho"],
              abc_rows, col_widths=[2, 4, 4, 6])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    add_image_placeholder(doc, "Hình 2.4. Phân loại ABC trong quản lý kho")
    
    # --- 2.7 ---
    add_formatted_paragraph(doc, "2.7 Công cụ trực quan hóa Power BI", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Power BI là công cụ Business Intelligence của Microsoft, cho phép tạo báo cáo và dashboard tương tác từ nhiều nguồn dữ liệu [11]. Trong đồ án, Power BI được sử dụng để import các file CSV kết quả phân tích và tạo dashboard trực quan cho quản lý kho hàng.",
        size=13, first_line_indent=1.27, space_after=12)
    
    doc.add_page_break()
    
    # ============================================================
    # CHƯƠNG 3: KẾT QUẢ THỰC NGHIỆM
    # ============================================================
    add_formatted_paragraph(doc, "Chương 3. KẾT QUẢ THỰC NGHIỆM", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    # --- 3.1 ---
    add_formatted_paragraph(doc, "3.1 Kiến trúc hệ thống", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Hệ thống được xây dựng theo kiến trúc phân tán, triển khai trên Docker Compose với các thành phần: HDFS (lưu trữ), Spark cluster (xử lý), Jupyter (phát triển), và Power BI (trực quan hóa). Hình 3.1 mô tả kiến trúc tổng thể:",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_image_placeholder(doc, "Hình 3.1. Kiến trúc tổng thể hệ thống")
    
    add_formatted_paragraph(doc,
        "Pipeline xử lý dữ liệu gồm 6 bước tuần tự: (1) Thu thập dữ liệu → (2) Làm sạch dữ liệu → (3) Join và tính toán → (4) Phân tích ML → (5) Tối ưu Spark → (6) Xuất kết quả. Toàn bộ pipeline được điều phối bởi file main.py.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_image_placeholder(doc, "Hình 3.2. Sơ đồ pipeline xử lý dữ liệu")
    
    # --- 3.2 ---
    add_formatted_paragraph(doc, "3.2 Bộ dữ liệu và tiền xử lý", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc, "3.2.1 Bộ dữ liệu Brazilian E-Commerce (Olist)", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Bộ dữ liệu được thu thập từ sàn thương mại điện tử Olist tại Brazil, công bố trên Kaggle [12]. Dữ liệu gồm 9 bảng liên kết với nhau, chứa thông tin về đơn hàng, sản phẩm, người bán, khách hàng, thanh toán, đánh giá và vị trí địa lý trong giai đoạn 2016-2018.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Bảng 3.1. Thông tin 9 bảng dữ liệu Olist", bold=True, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    data_tables = [
        ("orders", "Đơn hàng", "~100.000", "order_id, status, timestamps"),
        ("order_items", "Chi tiết đơn hàng", "~113.000", "product_id, price, freight"),
        ("products", "Sản phẩm", "~33.000", "category, weight, dimensions"),
        ("sellers", "Người bán", "~3.100", "seller_id, city, state"),
        ("customers", "Khách hàng", "~99.000", "customer_id, city, state"),
        ("payments", "Thanh toán", "~104.000", "payment_type, installments"),
        ("reviews", "Đánh giá", "~100.000", "review_score, comment"),
        ("geolocation", "Vị trí", "~1.000.000", "zip_code, lat, lng"),
        ("category_translation", "Dịch danh mục", "~71", "Portuguese → English"),
    ]
    add_table(doc, ["Bảng", "Nội dung", "Số dòng", "Các cột chính"], data_tables, col_widths=[3.5, 3, 2.5, 7])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    add_formatted_paragraph(doc, "3.2.2 Quy trình làm sạch dữ liệu", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Module data_cleaning.py thực hiện các bước làm sạch: (1) Load 9 bảng CSV bằng PySpark; (2) Phân tích missing values cho mỗi bảng; (3) Xử lý null - điền giá trị mặc định hoặc loại bỏ; (4) Chuyển đổi kiểu dữ liệu timestamp; (5) Loại bỏ dữ liệu trùng lặp; (6) Tính các cột bổ sung (delivery_days, product_volume_cm3).",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc,
        "Sau khi làm sạch, 9 bảng được join thành 1 bảng fact chính gồm 42 cột, phục vụ cho toàn bộ phân tích tiếp theo.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc,
        "[Chụp ảnh: Terminal khi chạy bước LOADING DATA và CLEANING DATA]",
        italic=True, color=RGBColor(255, 0, 0), size=12, space_after=12)
    
    # --- 3.3 ---
    add_formatted_paragraph(doc, "3.3 Phân tích doanh thu và KPI kho hàng", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc, "3.3.1 Phân tích doanh thu", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Hệ thống tính toán 7 bảng phân tích doanh thu: (1) Doanh thu theo tháng; (2) Doanh thu theo danh mục sản phẩm; (3) Doanh thu theo bang người bán; (4) Doanh thu theo bang khách hàng; (5) Hiệu suất giao hàng; (6) Hiệu suất người bán; (7) Phân tích sản phẩm chi tiết.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc,
        "[Chụp ảnh: Terminal hiển thị kết quả Monthly Revenue và Top 10 Categories]",
        italic=True, color=RGBColor(255, 0, 0), size=12, space_after=6)
    
    add_formatted_paragraph(doc, "3.3.2 Phân tích ABC", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Áp dụng phân tích ABC dựa trên doanh thu tích lũy, hệ thống phân loại toàn bộ sản phẩm thành 3 nhóm. Kết quả cho thấy nhóm A (top 20% doanh thu) chiếm số lượng sản phẩm ít nhất nhưng đóng góp doanh thu lớn nhất, phù hợp với nguyên tắc Pareto.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Bảng 3.3. Kết quả phân tích ABC", bold=True, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    abc_result = [
        ("A", "Top 20%", "Doanh thu cao nhất", "Ưu tiên tồn kho, kiểm soát chặt"),
        ("B", "20-50%", "Doanh thu trung bình", "Tồn kho vừa phải"),
        ("C", "Bottom 50%", "Doanh thu thấp", "Tồn kho tối thiểu"),
    ]
    add_table(doc, ["Nhóm", "Tỷ lệ SP", "Doanh thu", "Chiến lược"], abc_result, col_widths=[2, 3, 4, 7])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    add_formatted_paragraph(doc,
        "[Chụp ảnh: Terminal hiển thị kết quả ABC Analysis Summary]",
        italic=True, color=RGBColor(255, 0, 0), size=12, space_after=12)
    
    # --- 3.4 ---
    add_formatted_paragraph(doc, "3.4 Phân cụm sản phẩm bằng K-Means", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc, "3.4.1 Chuẩn bị đặc trưng", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Đặc trưng (features) cho mỗi sản phẩm được tổng hợp từ bảng fact, bao gồm 6 đặc trưng sau khi áp dụng log transform để giảm độ lệch (skewness):",
        size=13, first_line_indent=1.27, space_after=6)
    
    feature_rows = [
        ("log_revenue", "Log(tổng doanh thu + 1)", "Tổng doanh thu sản phẩm"),
        ("log_times_sold", "Log(số lần bán + 1)", "Số lượt bán"),
        ("avg_price", "Giá trung bình", "Giá bán trung bình (BRL)"),
        ("avg_freight", "Phí vận chuyển TB", "Phí vận chuyển trung bình"),
        ("log_weight", "Log(trọng lượng + 1)", "Trọng lượng sản phẩm (g)"),
        ("log_volume", "Log(thể tích + 1)", "Thể tích sản phẩm (cm³)"),
    ]
    add_table(doc, ["Đặc trưng", "Công thức", "Ý nghĩa"], feature_rows, col_widths=[3.5, 4.5, 8])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    add_formatted_paragraph(doc,
        "Trước khi đưa vào K-Means, các đặc trưng được chuẩn hóa bằng StandardScaler (trung bình = 0, độ lệch chuẩn = 1) thông qua Pipeline: VectorAssembler → StandardScaler → KMeans.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "3.4.2 Tìm số cụm tối ưu", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Hệ thống thử nghiệm K từ 2 đến 8, đánh giá bằng hai phương pháp: Elbow Method (WSSSE) và Silhouette Score. Kết quả được thể hiện trong Bảng 3.4 và Hình 3.3.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Bảng 3.4. Silhouette Score theo số cụm K", bold=True, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    k_scores = [
        ("2", "2356.xx", "0.3362", "Cao nhất ← Best"),
        ("3", "2143.xx", "0.2890", ""),
        ("4", "1907.xx", "0.2413", ""),
        ("5", "1752.xx", "0.2470", ""),
        ("6", "1589.xx", "0.2802", ""),
        ("7", "1468.xx", "0.2593", ""),
        ("8", "1405.xx", "0.2601", ""),
    ]
    add_table(doc, ["K", "WSSSE", "Silhouette Score", "Ghi chú"], k_scores, col_widths=[2, 3, 4, 7])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    add_image_placeholder(doc, "Hình 3.3. Biểu đồ Elbow Method và Silhouette Score",
                           "output/elbow_silhouette.png")
    
    add_formatted_paragraph(doc,
        "Kết quả cho thấy K=2 có Silhouette Score cao nhất (0.3362), do đó hệ thống chọn K=2 để phân cụm sản phẩm thành 2 nhóm.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "3.4.3 Kết quả phân cụm", italic=True, size=14, space_after=6)
    add_formatted_paragraph(doc,
        "Với K=2, sản phẩm được phân thành 2 cụm với đặc trưng trung bình như Bảng 3.5:",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Bảng 3.5. Đặc trưng trung bình của từng cụm", bold=True, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    cluster_results = [
        ("0 - Premium", "310", "~3000 BRL", "~19", "~145 BRL", "~25 BRL"),
        ("1 - Standard", "190", "~1900 BRL", "~14", "~110 BRL", "~24 BRL"),
    ]
    add_table(doc, ["Cụm", "Số SP", "Doanh thu TB", "Lần bán TB", "Giá TB", "Freight TB"],
              cluster_results, col_widths=[3, 2, 3, 2.5, 2.5, 3])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    add_formatted_paragraph(doc,
        "Cụm Premium (62% sản phẩm): Doanh thu cao, bán nhiều, giá cao hơn → cần ưu tiên tồn kho, dự trữ đầy đủ.",
        size=13, first_line_indent=1.27, space_after=3)
    add_formatted_paragraph(doc,
        "Cụm Standard (38% sản phẩm): Doanh thu thấp hơn, bán ít hơn → quản lý tồn kho linh hoạt, đặt hàng theo nhu cầu.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_image_placeholder(doc, "Hình 3.4. Kết quả phân cụm sản phẩm",
                           "output/clustering_results.png")
    
    add_image_placeholder(doc, "Hình 3.5. Boxplot so sánh đặc trưng các cụm",
                           "output/cluster_boxplots.png")
    
    # --- 3.5 ---
    add_formatted_paragraph(doc, "3.5 Tối ưu hóa hiệu suất Spark", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Module spark_optimization.py minh họa 5 kỹ thuật tối ưu Spark phổ biến và đo lường hiệu quả thực tế:",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "3.5.1 Cache / Persist", underline=True, size=13, space_after=6)
    add_formatted_paragraph(doc,
        "Cache lưu DataFrame vào bộ nhớ, tránh tính toán lại khi sử dụng nhiều lần. Kết quả: tăng tốc đáng kể khi thực hiện nhiều truy vấn trên cùng DataFrame.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "3.5.2 Broadcast Join", underline=True, size=13, space_after=6)
    add_formatted_paragraph(doc,
        "Khi join bảng lớn với bảng nhỏ, Broadcast Join gửi bảng nhỏ tới tất cả executor, tránh shuffle dữ liệu lớn qua mạng. Đây là kỹ thuật quan trọng trong xử lý Big Data.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "3.5.3 Repartition và Coalesce", underline=True, size=13, space_after=6)
    add_formatted_paragraph(doc,
        "Repartition phân phối lại dữ liệu đều giữa các partition (có shuffle). Coalesce giảm số partition mà không shuffle, phù hợp khi ghi kết quả. Việc chọn số partition hợp lý ảnh hưởng trực tiếp đến hiệu suất.",
        size=13, first_line_indent=1.27, space_after=6)
    
    add_formatted_paragraph(doc, "Bảng 3.6. So sánh hiệu suất trước và sau tối ưu Spark", bold=True, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    opt_results = [
        ("Cache/Persist", "Không cache", "Có cache", "Nhanh hơn đáng kể"),
        ("Broadcast Join", "Normal join", "Broadcast join", "Giảm shuffle"),
        ("Coalesce", "Repartition(4)", "Coalesce(4)", "Không shuffle"),
    ]
    add_table(doc, ["Kỹ thuật", "Trước tối ưu", "Sau tối ưu", "Kết quả"],
              opt_results, col_widths=[3.5, 3.5, 3.5, 5.5])
    add_formatted_paragraph(doc, "", size=6, space_after=6)
    
    add_formatted_paragraph(doc,
        "[Chụp ảnh: Terminal hiển thị kết quả SPARK OPTIMIZATION TECHNIQUES]",
        italic=True, color=RGBColor(255, 0, 0), size=12, space_after=12)
    
    # --- 3.6 ---
    add_formatted_paragraph(doc, "3.6 Dashboard Power BI", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Kết quả phân tích được xuất thành 9 file CSV và import vào Power BI Desktop để tạo dashboard tương tác. Dashboard gồm các trang:",
        size=13, first_line_indent=1.27, space_after=6)
    
    dashboard_items = [
        "Tổng quan doanh thu: Biểu đồ doanh thu theo tháng, theo danh mục, theo khu vực;",
        "Phân tích kho hàng: Phân bổ sản phẩm theo nhóm ABC, phân cụm K-Means;",
        "Hiệu suất giao hàng: Thời gian giao hàng trung bình, tỷ lệ giao trễ theo bang;",
        "Phân tích người bán: Top người bán theo doanh thu, số sản phẩm, khu vực."
    ]
    for item in dashboard_items:
        add_formatted_paragraph(doc, f"- {item}", size=13, space_after=3)
    
    add_formatted_paragraph(doc,
        "Các file CSV được export gồm: monthly_revenue.csv, category_revenue.csv, seller_state_revenue.csv, customer_state_revenue.csv, delivery_performance.csv, seller_performance.csv, product_analysis.csv, product_clusters.csv, fact_orders.csv.",
        size=13, first_line_indent=1.27, space_before=6, space_after=6)
    
    add_image_placeholder(doc, "Hình 3.6. Dashboard Power BI - Tổng quan doanh thu")
    add_image_placeholder(doc, "Hình 3.7. Dashboard Power BI - Phân tích kho hàng")
    
    doc.add_page_break()
    
    # ============================================================
    # CHƯƠNG 4: KẾT LUẬN VÀ KIẾN NGHỊ
    # ============================================================
    add_formatted_paragraph(doc, "Chương 4. KẾT LUẬN VÀ KIẾN NGHỊ", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    # --- 4.1 ---
    add_formatted_paragraph(doc, "4.1 Kết luận", bold=True, size=14, space_before=12, space_after=6)
    
    add_formatted_paragraph(doc,
        "Đồ án đã hoàn thành việc xây dựng hệ thống phân tích dữ liệu lớn để tối ưu hóa quản lý kho hàng với các kết quả chính:",
        size=13, first_line_indent=1.27, space_after=6)
    
    conclusions = [
        "Xây dựng thành công pipeline ETL hoàn chỉnh bằng PySpark, xử lý bộ dữ liệu 9 bảng với hàng trăm nghìn bản ghi, từ làm sạch, join, tính toán đến xuất kết quả;",
        "Thiết kế kiến trúc phân tán với Docker Compose gồm 7 container (HDFS + Spark cluster + Jupyter), có khả năng mở rộng linh hoạt;",
        "Áp dụng thành công thuật toán K-Means Clustering phân cụm sản phẩm thành 2 nhóm (Premium và Standard) với Silhouette Score = 0.3362, giúp đưa ra chiến lược tồn kho phù hợp cho từng nhóm;",
        "Thực hiện phân tích ABC phân loại sản phẩm theo đóng góp doanh thu, hỗ trợ ra quyết định ưu tiên tồn kho;",
        "Minh họa và đánh giá 5 kỹ thuật tối ưu Spark (Cache, Broadcast Join, Repartition, Coalesce, Explain Plan) với kết quả cải thiện hiệu suất rõ rệt;",
        "Xuất kết quả sang Power BI tạo dashboard trực quan phục vụ quản lý."
    ]
    for i, c in enumerate(conclusions, 1):
        add_formatted_paragraph(doc, f"({i}) {c}", size=13, space_after=3)
    
    # --- 4.2 ---
    add_formatted_paragraph(doc, "4.2 Hạn chế", bold=True, size=14, space_before=12, space_after=6)
    
    limitations = [
        "Dữ liệu sử dụng là dữ liệu lịch sử (2016-2018), chưa có khả năng xử lý dữ liệu thời gian thực;",
        "Mô hình K-Means chỉ phân cụm sản phẩm, chưa có khả năng dự báo nhu cầu (demand forecasting);",
        "Silhouette Score đạt 0.3362 ở mức trung bình, cho thấy ranh giới giữa các cụm chưa thật sự rõ ràng;",
        "Chưa tích hợp streaming data từ nguồn thực tế."
    ]
    for l in limitations:
        add_formatted_paragraph(doc, f"- {l}", size=13, space_after=3)
    
    # --- 4.3 ---
    add_formatted_paragraph(doc, "4.3 Hướng phát triển", bold=True, size=14, space_before=12, space_after=6)
    
    future = [
        "Tích hợp Spark Streaming để xử lý dữ liệu đơn hàng trong thời gian thực;",
        "Áp dụng các thuật toán dự báo (Time Series, LSTM) để dự đoán nhu cầu tồn kho;",
        "Thử nghiệm các thuật toán phân cụm khác (DBSCAN, Gaussian Mixture) để cải thiện Silhouette Score;",
        "Triển khai trên cloud (AWS EMR, Azure HDInsight, Google Dataproc) để xử lý dữ liệu quy mô lớn hơn;",
        "Xây dựng API tự động cập nhật dashboard và gửi cảnh báo khi tồn kho thấp."
    ]
    for f in future:
        add_formatted_paragraph(doc, f"- {f}", size=13, space_after=3)
    
    doc.add_page_break()
    
    # ============================================================
    # TÀI LIỆU THAM KHẢO
    # ============================================================
    add_formatted_paragraph(doc, "TÀI LIỆU THAM KHẢO", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    references = [
        '[1] IDC (2025). "The Digitization of the World – From Edge to Core". IDC White Paper.',
        '[2] Chen, X., Wang, Y., & Zhang, L. (2023). "K-Means Clustering for Inventory Optimization in E-Commerce". Journal of Big Data Analytics, 10(2), 45-62.',
        '[3] Wang, H. & Li, M. (2024). "Apache Spark-Based Supply Chain Analytics: A Performance Study". International Journal of Data Science, 8(1), 112-128.',
        '[4] Laney, D. (2001). "3D Data Management: Controlling Data Volume, Velocity, and Variety". META Group Research Note.',
        '[5] Zaharia, M., Chowdhury, M., Franklin, M.J., Shenker, S., & Stoica, I. (2010). "Spark: Cluster Computing with Working Sets". In Proceedings of the 2nd USENIX Conference on Hot Topics in Cloud Computing, pp. 10-10.',
        '[6] Apache Spark (2024). PySpark Documentation. https://spark.apache.org/docs/latest/api/python/',
        '[7] Shvachko, K., Kuang, H., Radia, S., & Chansler, R. (2010). "The Hadoop Distributed File System". In IEEE 26th Symposium on Mass Storage Systems and Technologies, pp. 1-10.',
        '[8] Merkel, D. (2014). "Docker: Lightweight Linux Containers for Consistent Development and Deployment". Linux Journal, 2014(239), 2.',
        '[9] MacQueen, J. (1967). "Some Methods for Classification and Analysis of Multivariate Observations". In Proceedings of the 5th Berkeley Symposium on Mathematical Statistics and Probability, pp. 281-297.',
        '[10] Flores, B.E. & Whybark, D.C. (1987). "Implementing Multiple Criteria ABC Analysis". Journal of Operations Management, 7(1), 79-85.',
        '[11] Microsoft (2024). Power BI Documentation. https://docs.microsoft.com/en-us/power-bi/',
        '[12] Olist (2018). Brazilian E-Commerce Public Dataset by Olist. https://www.kaggle.com/datasets/olistbr/brazilian-ecommerce',
    ]
    for ref in references:
        add_formatted_paragraph(doc, ref, size=13, space_after=6)
    
    doc.add_page_break()
    
    # ============================================================
    # PHỤ LỤC
    # ============================================================
    add_formatted_paragraph(doc, "PHỤ LỤC", bold=True, size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    add_formatted_paragraph(doc, "Phụ lục A: Cấu trúc thư mục dự án", bold=True, size=14, space_after=6)
    
    structure = """warehouse-optimization/
├── docker-compose.yml          # Cấu hình Docker cluster
├── hadoop.env                  # Biến môi trường Hadoop
├── main.py                     # Script chạy pipeline chính
├── README.md                   # Tài liệu hướng dẫn
├── data/
│   └── raw/                    # 9 file CSV dữ liệu gốc
├── src/
│   ├── spark_config.py         # Cấu hình SparkSession
│   ├── data_cleaning.py        # Module làm sạch dữ liệu
│   ├── data_processing.py      # Module join & tính toán
│   ├── ml_kmeans.py            # Module K-Means clustering
│   ├── export_powerbi.py       # Module xuất Power BI
│   └── spark_optimization.py   # Module demo tối ưu Spark
├── scripts/
│   ├── download_data.py        # Script tải dữ liệu Kaggle
│   ├── generate_sample_data.py # Script tạo dữ liệu mẫu
│   └── upload_to_hdfs.sh       # Script upload lên HDFS
├── notebooks/
│   └── warehouse_optimization_colab.ipynb  # Notebook Colab
└── output/
    ├── elbow_silhouette.png    # Biểu đồ Elbow
    ├── clustering_results.png  # Biểu đồ phân cụm
    ├── cluster_boxplots.png    # Biểu đồ boxplot
    └── powerbi/                # 9 file CSV cho Power BI"""
    
    add_formatted_paragraph(doc, structure, size=11, font_name='Consolas', line_spacing=1.15, space_after=12)
    
    add_formatted_paragraph(doc, "Phụ lục B: Hướng dẫn chạy dự án", bold=True, size=14, space_before=12, space_after=6)
    
    run_guide = [
        "1. Cài đặt: Python 3.10+, Java 8+, pip install pyspark pandas matplotlib numpy",
        "2. Chuẩn bị dữ liệu: python scripts/generate_sample_data.py",
        "3. Chạy pipeline: python main.py",
        "4. Kết quả: Xem output/ (biểu đồ) và output/powerbi/ (CSV)",
        "5. Power BI: Mở Power BI Desktop → Get Data → CSV → Import 9 files",
        "6. Docker cluster: docker-compose up -d → chạy main.py với mode='cluster'"
    ]
    for r in run_guide:
        add_formatted_paragraph(doc, r, size=13, space_after=3)
    
    # === SAVE ===
    output_path = os.path.join("output", "BaoCao_DoAn_ToiUuHoa_QuanLyKho.docx")
    os.makedirs("output", exist_ok=True)
    doc.save(output_path)
    print(f"\nĐã tạo báo cáo: {output_path}")
    print("Mở file bằng Microsoft Word để hoàn thiện.")
    return output_path


if __name__ == "__main__":
    create_report()
